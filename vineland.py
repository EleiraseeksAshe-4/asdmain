from docx import Document

from dataextractor import find_paras_between_headings, process_paragraphs


def vineland_ranking(standard_score):
    rankings = [{'ranking': 'High', 'subdomainlow': 21, 'subdomainhigh': 24, 'domainlow': 130, 'domainhigh': 140},
                {'ranking': 'Moderately High', 'subdomainlow': 18, 'subdomainhigh': 20, 'domainlow': 115,
                 'domainhigh': 129},
                {'ranking': 'Adequate', 'subdomainlow': 13, 'subdomainhigh': 17, 'domainlow': 86, 'domainhigh': 114},
                {'ranking': 'Moderately Low', 'subdomainlow': 10, 'subdomainhigh': 12, 'domainlow': 71,
                 'domainhigh': 85},
                {'ranking': 'Low', 'subdomainlow': 1, 'subdomainhigh': 9, 'domainlow': 20, 'domainhigh': 70}]

    for ranking in rankings:
        if standard_score >= ranking['domainlow'] and standard_score <= ranking['domainhigh']:
            return ranking['ranking']


def testvineland():
    print(vineland_ranking(135))
    print(vineland_ranking(116))
    print(vineland_ranking(87))
    print(vineland_ranking(71))
    print(vineland_ranking(70))

def find_table(doc, tableName:str) -> int:
    """
    Function to find the tables which need to be updated with value.
    Returns -1 if couldn't find table
    Args: int, string
    Returns: int
    """
    tableNum = -1
    for i, table in enumerate(doc.tables):
        if table.rows[0].cells[0].text == tableName:
            tableNum = i
            break
    return tableNum

def process_vineland_tables(vabs, doc):
    tableNames = ["ABC", "Subdomains", "Maladaptive Scale"]
    # print("HERE")
    # print(vabs["overall"])
    for tablen in tableNames:
        getTableIndex = find_table(doc,tablen)
        if getTableIndex != -1:
            for i, row in enumerate(doc.tables[getTableIndex].rows):
                if i == 0:
                    continue
                temp = {}
                for j, cell in enumerate(row.cells):
                    if j == 0:
                        continue
                    if doc.tables[getTableIndex].cell(i,0).text == doc.tables[getTableIndex].cell(i,1).text:
                        break
                    if tablen == "ABC":
                        if doc.tables[getTableIndex].cell(0,j).text == "Standard Score (SS)":
                            if doc.tables[getTableIndex].cell(i,0).text == "Adaptive Behavior Composite":
                                vabs['overall_std'] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Communication":
                                vabs["comms_std"] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Daily Living Skills":
                                vabs["daily_std"] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Socialization":
                                vabs["socialization_std"] = cell.text
                        elif doc.tables[getTableIndex].cell(0,j).text == "90% Confidence Interval":
                            if doc.tables[getTableIndex].cell(i,0).text == "Adaptive Behavior Composite":
                                vabs['overall_ci'] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Communication":
                                vabs["comms_ci"] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Daily Living Skills":
                                vabs["daily_ci"] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Socialization":
                                vabs["socialization_ci"] = cell.text
                        elif doc.tables[getTableIndex].cell(0,j).text == "Percentile Rank":
                            if doc.tables[getTableIndex].cell(i,0).text == "Adaptive Behavior Composite":
                                vabs['overall_pr'] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Communication":
                                vabs["comms_pr"] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Daily Living Skills":
                                vabs["daily_pr"] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Socialization":
                                vabs["socialization_pr"] = cell.text
                        elif doc.tables[getTableIndex].cell(0,j).text == "SS Minus Mean SS*":
                            if doc.tables[getTableIndex].cell(i,0).text == "Adaptive Behavior Composite":
                                vabs['overall_ssmmss'] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Communication":
                                vabs["comms_ssmmss"] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Daily Living Skills":
                                vabs["daily_ssmmss"] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Socialization":
                                vabs["socialization_ssmmss"] = cell.text
                        elif doc.tables[getTableIndex].cell(0,j).text == "Strength or Weakness**":
                            if doc.tables[getTableIndex].cell(i,0).text == "Adaptive Behavior Composite":
                                vabs['overall_sw'] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Communication":
                                vabs["comms_sw"] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Daily Living Skills":
                                vabs["daily_sw"] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Socialization":
                                vabs["socialization_sw"] = cell.text
                        else:
                            if doc.tables[getTableIndex].cell(i,0).text == "Adaptive Behavior Composite":
                                vabs['overall_br'] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Communication":
                                vabs["comms_br"] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Daily Living Skills":
                                vabs["daily_br"] = cell.text
                            elif doc.tables[getTableIndex].cell(i,0).text == "Socialization":
                                vabs["socialization_br"] = cell.text
                    elif tablen == "Subdomains":
                        if doc.tables[getTableIndex].cell(0,j).text == "Raw Score":
                            tempHeader = doc.tables[getTableIndex].cell(i,0).text.replace(" ", "_").lower() + "_rs"
                            vabs[tempHeader] = cell.text
                        elif doc.tables[getTableIndex].cell(0,j).text == "v-Scale Score ( vS)":
                            tempHeader = doc.tables[getTableIndex].cell(i,0).text.replace(" ", "_").lower() + "_vss"
                            vabs[tempHeader] = cell.text
                        elif doc.tables[getTableIndex].cell(0,j).text == "Age Equivalent":
                            tempHeader = doc.tables[getTableIndex].cell(i,0).text.replace(" ", "_").lower() + "_ae"
                            vabs[tempHeader] = cell.text
                        elif doc.tables[getTableIndex].cell(0,j).text == "Growth Scale Value":
                            tempHeader = doc.tables[getTableIndex].cell(i,0).text.replace(" ", "_").lower() + "_gsv"
                            vabs[tempHeader] = cell.text
                        elif doc.tables[getTableIndex].cell(0,j).text == "Percent Estimated":
                            tempHeader = doc.tables[getTableIndex].cell(i,0).text.replace(" ", "_").lower() + "_pe"
                            vabs[tempHeader] = cell.text
                        elif doc.tables[getTableIndex].cell(0,j).text == "vS Minus Mean vS*":
                            tempHeader = doc.tables[getTableIndex].cell(i,0).text.replace(" ", "_").lower() + "_vsmmvs"
                            vabs[tempHeader] = cell.text
                        elif doc.tables[getTableIndex].cell(0,j).text == "vS Minus Mean vS*":
                            tempHeader = doc.tables[getTableIndex].cell(i,0).text.replace(" ", "_").lower() + "_vsmmvs"
                            vabs[tempHeader] = cell.text
                        elif doc.tables[getTableIndex].cell(0,j).text == "Strength or Weakness**":
                            tempHeader = doc.tables[getTableIndex].cell(i,0).text.replace(" ", "_").lower() + "_sow"
                            vabs[tempHeader] = cell.text
                        else:
                            tempHeader = doc.tables[getTableIndex].cell(i,0).text.replace(" ", "_").lower() + "_br"
                            vabs[tempHeader] = cell.text
                    else:
                        if doc.tables[getTableIndex].cell(0,j).text == "Raw Score":
                            tempHeader = doc.tables[getTableIndex].cell(i,0).text.replace(" ", "_").lower() + "_rs"
                            
                            vabs[tempHeader] = cell.text
                        else:
                            tempHeader = doc.tables[getTableIndex].cell(i,0).text.replace(" ", "_").lower() + "_vss"
                            vabs[tempHeader] = cell.text
                    # print(doc.tables[getTableIndex].cell(0,j).text)
                    # print(cell.text)
                    temp[doc.tables[getTableIndex].cell(0,j).text] = cell.text
                vabs[doc.tables[getTableIndex].cell(i,0).text] = temp
    return vabs

def process_vineland(doc):
    vabs = {}

    keypoints = [('overall', 'Adaptive Behavior', 'Communication Domain', False),
                 ('comms', 'Communication Domain', 'Daily Living Skills Domain', False),
                 ('daily', 'Daily Living Skills Domain', 'Socialization Domain', False),
                 ('socialization', 'Socialization Domain', 'Motor Skills Domain', False),
                 ('motor', 'Motor Skills Domain',
                  'Domain-Level Strengths/Weaknesses and Pairwise Difference Comparisons', True),
                 ('subdomain', 'Domain-Level Strengths/Weaknesses and Pairwise Difference Comparisons',
                  'Maladaptive Behavior', True),
                 ('maladaptive', 'Maladaptive Behavior', 'INTERVENTION GUIDANCE', False)]

    # heading = keypoints[0]
    for heading in keypoints:
        paras = find_paras_between_headings(doc, 0, heading[1], heading[2], heading[3])
        vabs[heading[0]] = paras
        print(vabs[heading[0]])

    return vabs


def fetch_vabs_data(context):
    # vabs = {
    #     'overall_std': 68,
    #     'overall_rating': "Low",
    #     'comms_std': 79,
    #     'comms_rating': "Moderately Low",
    #     'daily_std': 71,
    #     'daily_rating': "Moderately Low",
    #     'socialization_std': 54,
    #     'socialization_rating': "Low"
    # }

    vabs = {}
    active_path = context['active_path']
    vineland_doc_path = f"{active_path}/word/vineland.docx"
    vineland_doc = Document(vineland_doc_path)
    vabs_text = process_vineland(vineland_doc)
    process_paragraphs(vabs, vabs_text, context['style'])

    # print("Process Paragraph")
    
    process_vineland_tables(vabs, vineland_doc)

    # tableNames = ["ABC", "Subdomains"]
    # for tablen in tableNames:
    #     getTableIndex = find_table(vineland_doc,tablen)
    #     for i, row in enumerate(vineland_doc.tables[getTableIndex].rows):
            
    #         for j, cell in enumerate(row.cells): 

    ratingscales = context['rs']
    ratingscales['vabs'] = vabs
    # print(vabs)
    return vabs
