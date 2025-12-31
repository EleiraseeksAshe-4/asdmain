from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import csv

from docx.shared import Pt
from docxtpl import RichText

def add_paragraph_to_rich_text(rich_text, paragraph, style_override=None):
    for run in paragraph.runs:
        if style_override == None:
            rich_text += run.text
            # rich_text.add(
            #     run.text,
            #     bold=run.bold,
            #     italic=run.italic,
            #     underline=run.underline,
            #     color=run.font.color.rgb if run.font.color else None,
            #     size=Pt(12),#run.font.size.pt if run.font.size else None,
            #     font='Times New Roman')#run.font.name)
        else:
            rich_text += run.text
            # rich_text.add(
            #     run.text,
            #     bold=run.bold,
            #     italic=run.italic,
            #     underline=run.underline,
            #     color=style_override.font.color.rgb if style_override.font.color else None,
            #     size=style_override.font.size.pt if style_override.font.size else None,
            #     font=style_override.font.name)

    #    rich_text.add('\n')
    rich_text += '\n'
    return rich_text

def convert_paragraphs_to_rich_text(paragraphs, style_override=None):
    rich_text = ''
    for paragraph in paragraphs:
        rich_text = add_paragraph_to_rich_text(rich_text, paragraph, style_override)
    return rich_text

def process_paragraphs(ratingscale, text, style=None):

    for key in text.keys():
        try:
            rt = convert_paragraphs_to_rich_text(text[key], style)
            ratingscale[key] = rt
        except TypeError:
            print('Type error')
            pass
        except IndexError:
            print('Index error')
            pass
        finally:
            pass


def check_doc_content(eltomatch):
    # Iterate through the body elements of the document
    for element in eltomatch.part.document.element.body:
        # Check if the element is a paragraph
        if element == eltomatch:
            paragraph = element._element.getparent()
            print(paragraph.text)

        # Check if the element is a table
        elif element.tag.endswith('tbl'):
            table = element._element.getparent()
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        print(paragraph.text)


def copy_tables(input_doc, output_doc, skipCol):
    # Iterate over each table in the input document
    for i, table in enumerate(input_doc.tables):
        # Add a paragraph before the table with the index
        output_doc.add_paragraph(f"Table {i + 1}:")

        # Clone the table and add it to the output document
        skip1stCol = skipCol
        if i == 0:
            skip1stCol = False
        clone_table(table, output_doc, skip1stCol)


def clone_table(table, output_doc, skip_first_column):
    # Create a new table in the output document with the same dimensions
    columns = len(table.columns)
    if skip_first_column:
        columns = columns - 1

    new_table = output_doc.add_table(rows=0, cols=columns)

    # Copy the style of the table (if any)
    new_table.style = table.style
    new_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Iterate over the rows in the input table
    for row in table.rows:
        # Add a new row at the end of the new table
        new_row = new_table.add_row()
        start_index = 1 if skip_first_column else 0
        for idx in range(start_index, len(table.columns)):
            # Copy the text from the input cell to the new cell
            newcell = new_row.cells[idx - start_index]
            style_cell(newcell)
            newcell.text = row.cells[idx].text


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000"},
        bottom={"sz": 12, "color": "#00FF00", "space": "0"}
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def style_cell(cell):
    set_cell_border(cell, bottom={"sz": 12, "color": "#000000", "val": "single"})
    # top={"sz": ("w", "24"), "color": ("w", "auto"), "val": ("w", "single")},

    # start={"sz": ("w", "24"), "color": ("w", "auto"), "val": ("w", "single")},
    # end={"sz": ("w", "24"), "color": ("w", "auto"), "val": ("w", "single")})


def extract_and_write_section(input_doc, start_title, end_title):
    # Flag to indicate if we are within the desired section
    capture = False

    extracted_paragraphs = []
    # Iterate through each paragraph in the input document
    for paragraph in input_doc.paragraphs:
        paratext = paragraph.text.strip()
        print(paratext)
        if paratext.startswith(start_title):
            # Start capturing the content if the start title is found
            capture = True

        if capture:
            # Copy paragraph and its style to the output document
            extracted_paragraphs.append(paragraph._element)

            # copied_paragraph.add_run(paragraph.text).style = paragraph.style

        if paratext.startswith(end_title):
            # Stop capturing after the end title paragraph is included
            capture = False
            break

    temp_doc = Document()
    fragment_content = temp_doc.add_paragraph()

    for paragraph in extracted_paragraphs:
        fragment_content.add_run(paragraph)

    return fragment_content


def list_allparagraphs(doc):
    paras = []

    # Iterate through each paragraph to check its style
    for index, paragraph in enumerate(doc.paragraphs):
        paras.append((index, paragraph.style.name, paragraph.text))

    return paras

def find_headings(doc):
    # List to hold headings
    headings = []

    # Iterate through each paragraph to check its style
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            headings.append(paragraph.text, paragraph)

    return headings

def find_paras_between_headings(doc, startidx, starttext, endtext, check_heading):
    # List to hold headings
    paras = []

    # Iterate through each paragraph to check its style
    index = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        paratext = paragraph.text.strip()
        print(paratext)
        textmatch = paratext.startswith(starttext)
        if (paragraph.style.name.startswith('Heading') or check_heading) and textmatch:
            index = idx + 1
            break

    for paragraph in doc.paragraphs[index:]:
        style_name = paragraph.style.name
        para_text = paragraph.text.strip()
        if (style_name.startswith('Heading') or check_heading) and para_text.strip().startswith(endtext):
            break
        else:
            paras.append(paragraph)

    return paras



# def create_template_doc():
# Create header - Logo, Contact Text Box, Psychologist Text Box
# Table of Contents
# Confidential Neuropsychological Evaluation  (Table)
# Name
# Date of Birth
# Age
# Sex
# Evaluators Names and Titles
# Referral Information
# Referred by
# Reasons for Concern
# Summary and Impressions
# Background Information
# Developmental History
# Birth
# Milestones
# Motor
# Language
# Social
# Educational History
# Medical History
# Treatment History
# Previous Evaluations
# Family Information
# Medical
# Psychological
# Drug use
# Behavioural Observations
# Family
# Clinician
# Assessment Procedures and Instruments
# < 5 CARS2, GARS3, VABS
# > 5 CARS2, GARS3, VABS, CEFI, ADHDT2,
# Test Results and Interpretation
# Diagnosis
# DSM-5 or ICD-10
# Codes for Diagnosis
# Recommendations
# Therapeutic
# Educational
# Family Support
# Resources
# Support Groups
# Followup
# Summary and Conclusions
# Appendices
# Supporting Documents
# Consent and Confidentiality
# Create footer - Patient Name, Date of Birth, Age, Page Number


# Specify the path to your documents
# fileroot = '/Users/johnbridges/Dropbox/NewHope Psychology/JoyConsulting/GateInc/ExpertReports/'
# client = 'ChrisMacArthurSettle'
# gars = 'narrativereport.docx'
# cefi = '2024.4.19 ChrMcA CEFI.docx'
# vineland = '2024.4.19 ChrMcA Vineland.docx'
# cars2 = ''
# input_file_path = fileroot + client + '/word/' + vineland
# output_file_path = fileroot + client + '/word/' + 'extracted_' + vineland

#CARS2

# # Vineland
# start_title = 'OVERALL SUMMARY'
# end_title = 'REPORT TO PARENT'
# skipCols = False

# CEFI
# start_title = 'About'
# end_title = 'References'
# skipCols = False

# GARS3
# start_title = 'Composite Results'
# end_title = 'Maladaptive Speech (MS)'
# skipCols = True

# Load the input document
# input_doc = Document(input_file_path)
# vabs = process_vineland(input_doc)
#
# print(vabs)
# print("Headings ...")
# paras = list_allparagraphs(input_doc)
# csvpath = fileroot + client + '/word/' + 'vineland.csv'
# with open(csvpath, 'w', newline='') as file:
#     writer = csv.writer(file)
#     writer.writerows(paras)
#
# print("Done with Headings")
# Create a new document for output
# output_doc = Document()

# Call the function
#extract_and_write_section(input_doc, output_doc, start_title, end_title)
#copy_tables(input_doc, output_doc, skipCols)

# Save the output document
# output_doc.save(output_file_path)
