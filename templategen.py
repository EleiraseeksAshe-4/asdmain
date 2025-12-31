import datetime
from io import BytesIO

import jinja2
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.shared import Pt 
from docx.opc.oxml import qn
from docx.oxml import OxmlElement
from docxtpl import DocxTemplate, RichText, Subdoc
from jinja2 import Undefined, Environment

from adhdt2 import fetch_adhdt2_data
from cars import fetch_cars2_data
from cefi import  fetch_cefi_data
from gars import fetch_gars_data
from vineland import  fetch_vabs_data
from vineland import find_table

class SilentUndefined(Undefined):
    def _fail_with_undefined_error(self, *args, **kwargs):
        return ''
    def __str__(self):
        return ''
    def __getattr__(self, name):
        return self

def insert_ordinals(data):

    print(data)
    suffix = "_ord"
    new_items = {}
    for key, value in list(data.items()):
        if isinstance(value, dict):
            insert_ordinals(value)
        else:
            if "percent" in key:
                ordinal = get_ordinal(value)
                newkey = f"{key}{suffix}"
                new_items[newkey] = ordinal

    if len(new_items) > 0:
        data.update(new_items)

    return data

def get_ordinal(percentage):
    # Convert percentage to integer
    if percentage == None:
        percentage = 0
    print(percentage)
    num = int(float(percentage))

    # Handle special cases for 11th, 12th, and 13th
    if 10 <= num % 100 <= 20:
        suffix = 'th'
    else:
        # Determine suffix based on last digit
        last_digit = num % 10
        if last_digit == 1:
            suffix = 'st'
        elif last_digit == 2:
            suffix = 'nd'
        elif last_digit == 3:
            suffix = 'rd'
        else:
            suffix = 'th'

    return f"{suffix}"



def calculate_age(birthdate_str):
    # Convert the string to a datetime object
    birthdate = datetime.datetime.strptime(birthdate_str, "%m/%d/%Y").date()

    # Get today's date
    today = datetime.date.today()

    # Calculate the difference in years and months
    difference = relativedelta(today, birthdate)
    years = difference.years
    months = difference.months
    days = difference.days

    # Calculate decimal age in years
    days_in_year = 365.25  # average considering leap years
    decimal_age = years + (months * 30 + days) / days_in_year

    # Create a formatted string for years and months
    age_str = f"{years} years {months} months"
    decimal_age_str = f"{decimal_age:.1f}"

    return years, months, decimal_age_str

def paragraphs_to_text(paragraphs):
    result = ''
    for value in paragraphs:
        for run in value.runs:
            result += run.text
        result += '\n'
    return result

def format_date_longform(date):
    # Extract day, month, and year from the date
    day = date.day
    month = date.strftime("%B")
    year = date.year

    # Get the day suffix
    suffix = get_ordinal(day)

    # Format the date in the long form
    formatted_date = f"{day}{suffix} {month} {year}"
    return formatted_date

def format_date_longform(date):
    # Extract day, month, and year from the date
    day = date.day
    month = date.strftime("%B")
    year = date.year

    # Get the day suffix
    suffix = get_ordinal(day)

    return day, suffix, month, year

def add_superscript(run, text):
    run.add_text(text)
    r = run._r
    rPr = r.get_or_add_rPr()
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'superscript')
    rPr.append(vertAlign)

def create_richtext_with_superscript(date):
    day, suffix, month, year = format_date_longform(date)
    rt = RichText()
    rt.add(day)
    rt.add(suffix, superscript=True)
    rt.add(f" {month} {year}")
    return rt

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

templates = {}

# Can probably make this generic by scanning a directory and loading all the templates
def load_templates(fileroot):
    templates.update({'master' : DocxTemplate(f"{fileroot}/templates/GATEwaymaster_template.docx")})
    master = templates['master']
    master.render_init()
    templates.update({'gars' : master.new_subdoc(f"{fileroot}/templates/gars3_template.docx")})
    templates.update({'vabs' : master.new_subdoc(f"{fileroot}/templates/vineland_template.docx")})
    templates.update({'cefi' : master.new_subdoc(f"{fileroot}/templates/cefi_template.docx")})
    templates.update({'cars2' : master.new_subdoc(f"{fileroot}/templates/cars2_template.docx")})
    # templates.update({'adhdt2' : master.new_subdoc(f"{fileroot}/templates/adhdt2_template.docx")})
    templates.update({'npscores' : master.new_subdoc(f"{fileroot}/templates/npscores_template.docx")})

    # for key, template in templates.items():
    #     template.render_init()


jinja_env = Environment(undefined=SilentUndefined)
# fileroot = '/Users/johnbridges/Dropbox/NewHope Psychology/JoyConsulting/GateInc/ExpertReports/'
fileroot = 'resources/'
load_templates(fileroot)

style = None


clientref = 'PaxOls'

active_path = f"{fileroot}clients/active/{clientref}"

todaysdate = datetime.date.today()
todays_date_formatted = todaysdate.strftime("%m/%d/%Y")
todays_date_filefriendly = todaysdate.strftime("%m_%d_%Y")

pronoun = 'She'
clientfullname = "Paxtyn Olson"
clientfirstname = clientfullname.split()[0]

clientdob = '03/17/2011'
ageyears, agemonths, decimal_age_str = calculate_age(clientdob)

ratingscales = {}

context = { 'examiner': "Dr. Sherrie L. Baehr",
            'fileroot': fileroot,
            'active_path': active_path,
            'clientref': clientref,
            'pronoun': pronoun,
            'clientfullname': clientfullname,
            'client_firstname': clientfullname.split()[0],
            'clientdob': clientdob,
            'testingdate': "04/28/2024",
            'ageyears': ageyears,
            'agemonths': agemonths,
            'clientage': decimal_age_str,
            'reportdate': todays_date_formatted,
            'longformdate': create_richtext_with_superscript(todaysdate),
            'style' : None,
            'rs': ratingscales
            }



fetch_gars_data(context)
fetch_vabs_data(context)
fetch_cefi_data(context)
fetch_cars2_data(context)
# fetch_adhdt2_data(context)

# print(fetch_vabs_data(context))

insert_ordinals(ratingscales)


doc = templates['master']

subdoc_context = {}
for key, template in templates.items():
    if key == 'vabs':
        subdoc = templates[key]
        ### Update tables
        # print(context["rs"]["vabs"]["overall"])
        tableNames = ["ABC", "Subdomains"]
        for tablen in tableNames:
            getTableIndex = find_table(subdoc,tablen)
            for i, row in enumerate(subdoc.tables[getTableIndex].rows):
                if i == 0:
                    continue
                for j, cell in enumerate(row.cells):
                    if j == 0:
                        continue

                    if cell.text == subdoc.tables[getTableIndex].cell(i,0).text:
                        break
                    # print(cell.text)
                    # print(context["rs"]["vabs"][subdoc.tables[getTableIndex].cell(i,0).text])
                    # print(subdoc.tables[getTableIndex].cell(0,j).text)
                    try:
                        cell.text = context["rs"]["vabs"][subdoc.tables[getTableIndex].cell(i,0).text][subdoc.tables[getTableIndex].cell(0,j).text]
                    except:
                        cell.text = ""
                    for paragraph in cell.paragraphs:
                        paragraph.style = subdoc.styles['Table Paragraph']
                        paragraph.alignment = 1 # Center alignement
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
        ## CLEANUP TABLES
        for tablen in tableNames:
            getTableIndex = find_table(subdoc,tablen)
            i = 0
            while i < len(subdoc.tables[getTableIndex].rows):
                if subdoc.tables[getTableIndex].cell(i,1).text == "":
                    tempTable = subdoc.tables[getTableIndex]
                    row = subdoc.tables[getTableIndex].rows[i]
                    remove_row(tempTable,row)
                    i = 0
                else:
                    i+=1
        ## Adaptive Behaviour Area Table fill
        getTableIndex = find_table(subdoc,"Adaptive Behaviour Area")
        for i, row in enumerate(subdoc.tables[getTableIndex].rows):
            if i == 0:
                continue
            for j, cell in enumerate(row.cells):
                if j == 0:
                    continue
                if j == 1:
                    if subdoc.tables[getTableIndex].cell(i,0).text == "Overall Summary Score":
                        cell.text = context["rs"]["vabs"]['overall_std']
                    elif subdoc.tables[getTableIndex].cell(i,0).text == "Communication Skills":
                        cell.text = context["rs"]["vabs"]['comms_std']
                    elif subdoc.tables[getTableIndex].cell(i,0).text == "Daily Living Skills":
                        cell.text = context["rs"]["vabs"]['daily_std']
                    elif subdoc.tables[getTableIndex].cell(i,0).text == "Socialization Skills":
                        cell.text = context["rs"]["vabs"]['socialization_std']
                    else:
                        break
        
        i = 0
        getTableIndex = find_table(subdoc,"Adaptive Behaviour Area")
        while i < len(subdoc.tables[getTableIndex].rows):
            if subdoc.tables[getTableIndex].cell(i,1).text == "":
                tempTable = subdoc.tables[getTableIndex]
                row = subdoc.tables[getTableIndex].rows[i]
                remove_row(tempTable,row)
                i = 0
            else:
                i+=1

        getTableIndex = find_table(subdoc,"Maladaptive Scale")
        for i, row in enumerate(subdoc.tables[getTableIndex].rows):
            if i == 0:
                continue
            for j, cell in enumerate(row.cells):
                if j == 0:
                    continue
                if j == 1:
                    if subdoc.tables[getTableIndex].cell(i,0).text == "Internalizing":
                        cell.text = context["rs"]["vabs"]['internalizing_rs']
                    elif subdoc.tables[getTableIndex].cell(i,0).text == "Externalizing":
                        cell.text = context["rs"]["vabs"]['externalizing_rs']
                    else:
                        break
        #         # else:
        #         #     if subdoc.tables[getTableIndex].cell(i,0).text == "Overall Summary Score":
        #         #         cell.text = context["rs"]["vabs"]['overall_rating']
        #         #     elif subdoc.tables[getTableIndex].cell(i,0).text == "Communication Skills":
        #         #         cell.text = context["rs"]["vabs"]['comms_rating']
        #         #     elif subdoc.tables[getTableIndex].cell(i,0).text == "Daily Living Skills":
        #         #         cell.text = context["rs"]["vabs"]['daily_rating']
        #         #     elif subdoc.tables[getTableIndex].cell(i,0).text == "Socialization Skills":
        #         #         cell.text = context["rs"]["vabs"]['socialization_rating']
        #         #     else:
        #         #         break
        #         # for paragraph in cell.paragraphs:
        #         #     paragraph.style = subdoc.styles['Table Paragraph']
        #         #     paragraph.alignment = 1 # Center alignement
        #         #     for run in paragraph.runs:
        #         #         run.font.size = Pt(8)
        # print(context["rs"]["vabs"]['overall_std'])
        print(context["rs"]["vabs"].keys())
        i = 0
        getTableIndex = find_table(subdoc,"Maladaptive Scale")
        while i < len(subdoc.tables[getTableIndex].rows):
            if subdoc.tables[getTableIndex].cell(i,1).text == "":
                tempTable = subdoc.tables[getTableIndex]
                row = subdoc.tables[getTableIndex].rows[i]
                remove_row(tempTable,row)
                i = 0
            else:
                i+=1
        context[f"{key}_doc"] = subdoc
    if key != 'master':
        subdoc = templates[key]
        context[f"{key}_doc"] = subdoc

        # insert_docx_content(doc, subdoc, f"{key}_doc")
# print(context["rs"])



doc.render(context, jinja_env=jinja_env)



doc_io = BytesIO()
doc.save(doc_io)

doc.save("tmp/fred.docx")

final_docxtplt = DocxTemplate(doc_io)

final_docxtplt.render(context, jinja_env=jinja_env)



output_file_path = active_path + '/word/'+ clientref + '_GATEway Neuropsychological Evaluation_' + todays_date_filefriendly + '.docx'
print(output_file_path)
final_docxtplt.save(output_file_path)
